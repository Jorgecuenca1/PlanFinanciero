"""
Script para crear manuales de usuario en Word
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT

def crear_estilo_titulo(doc):
    """Configura estilos del documento"""
    pass

def agregar_titulo_principal(doc, texto):
    """Agrega título principal centrado"""
    titulo = doc.add_heading(texto, 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return titulo

def agregar_seccion(doc, texto):
    """Agrega título de sección"""
    return doc.add_heading(texto, 1)

def agregar_subseccion(doc, texto):
    """Agrega subtítulo"""
    return doc.add_heading(texto, 2)

def agregar_tabla_credenciales(doc, email, password, rol):
    """Agrega tabla con credenciales"""
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'

    celdas = [
        ('Campo', 'Valor'),
        ('Email', email),
        ('Contraseña', password),
        ('Rol', rol),
    ]

    for i, (campo, valor) in enumerate(celdas):
        row = table.rows[i]
        row.cells[0].text = campo
        row.cells[1].text = valor
        if i == 0:
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].paragraphs[0].runs[0].bold = True

    return table

def agregar_lista(doc, items):
    """Agrega lista con viñetas"""
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
    return

def agregar_nota_importante(doc, texto):
    """Agrega nota destacada"""
    p = doc.add_paragraph()
    run = p.add_run('⚠️ IMPORTANTE: ')
    run.bold = True
    p.add_run(texto)
    return p

def agregar_modulos_proximos(doc):
    """Agrega sección de módulos próximos"""
    agregar_seccion(doc, '🚀 Módulos Próximamente Disponibles')

    doc.add_paragraph(
        'El Sistema de Habilitación de Servicios de Salud se encuentra en constante '
        'evolución. Próximamente se integrarán los siguientes módulos para complementar '
        'la gestión integral de calidad en su institución:'
    )

    # SIAU
    agregar_subseccion(doc, '📋 SIAU - Sistema de Información y Atención al Usuario')
    doc.add_paragraph(
        'Módulo para la gestión integral de peticiones, quejas, reclamos, sugerencias '
        'y felicitaciones (PQRSF) de los usuarios. Incluirá:'
    )
    agregar_lista(doc, [
        'Registro y seguimiento de PQRSF',
        'Tiempos de respuesta según normatividad',
        'Reportes estadísticos de satisfacción',
        'Encuestas de satisfacción al usuario',
        'Indicadores de calidad en atención'
    ])

    # RIAS
    agregar_subseccion(doc, '🏥 RIAS - Rutas Integrales de Atención en Salud')
    doc.add_paragraph(
        'Módulo para el seguimiento y gestión de las Rutas Integrales de Atención '
        'en Salud según la Resolución 3280 de 2018. Incluirá:'
    )
    agregar_lista(doc, [
        'Gestión de rutas de promoción y mantenimiento',
        'Rutas de grupos de riesgo',
        'Rutas específicas de atención',
        'Seguimiento de intervenciones',
        'Indicadores de cobertura y cumplimiento'
    ])

    # PAMEC
    agregar_subseccion(doc, '✅ PAMEC - Programa de Auditoría para el Mejoramiento de la Calidad')
    doc.add_paragraph(
        'Módulo para la implementación del Programa de Auditoría para el Mejoramiento '
        'de la Calidad en Salud. Incluirá:'
    )
    agregar_lista(doc, [
        'Autoevaluación institucional',
        'Planes de mejoramiento',
        'Seguimiento a acciones correctivas',
        'Auditoría de historias clínicas',
        'Indicadores de calidad SOGC',
        'Comités de calidad'
    ])

    doc.add_paragraph()
    agregar_nota_importante(doc,
        'Estos módulos estarán disponibles en próximas actualizaciones del sistema. '
        'Se notificará oportunamente cuando estén habilitados.'
    )

# ============================================================
# MANUAL ADMINISTRADOR
# ============================================================
def crear_manual_admin():
    doc = Document()

    # Título
    agregar_titulo_principal(doc, 'MANUAL DE USUARIO\nADMINISTRADOR DEL SISTEMA')

    doc.add_paragraph()
    p = doc.add_paragraph('Sistema de Habilitación de Servicios de Salud')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph('Resolución 3100 de 2019')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    # Credenciales
    agregar_seccion(doc, '1. Credenciales de Acceso')
    doc.add_paragraph('Utilice las siguientes credenciales para ingresar al sistema:')
    doc.add_paragraph()
    agregar_tabla_credenciales(doc, 'admin@habilitacion.com', 'Admin123!', 'Super Administrador')
    doc.add_paragraph()
    doc.add_paragraph('URL de acceso: http://127.0.0.1:8012/')

    # Descripción del rol
    agregar_seccion(doc, '2. Descripción del Rol')
    doc.add_paragraph(
        'Como Super Administrador del sistema, usted tiene acceso completo a todas las '
        'funcionalidades. Este rol está diseñado para la gestión global del sistema, '
        'permitiendo administrar múltiples entidades prestadoras de servicios de salud.'
    )

    # Funcionalidades
    agregar_seccion(doc, '3. Funcionalidades Disponibles')

    agregar_subseccion(doc, '3.1 Dashboard')
    doc.add_paragraph('Panel principal con estadísticas generales del sistema:')
    agregar_lista(doc, [
        'Total de entidades registradas',
        'Total de sedes activas',
        'Estadísticas de evaluaciones',
        'Indicadores de cumplimiento global'
    ])

    agregar_subseccion(doc, '3.2 Gestión de Entidades')
    doc.add_paragraph('Administración completa de entidades prestadoras:')
    agregar_lista(doc, [
        'Crear nuevas entidades prestadoras',
        'Editar información de entidades existentes',
        'Visualizar todas las entidades del sistema',
        'Gestionar sedes de cada entidad',
        'Crear usuarios automáticamente al registrar entidades'
    ])

    agregar_subseccion(doc, '3.3 Gestión de Usuarios')
    doc.add_paragraph('Control de acceso al sistema:')
    agregar_lista(doc, [
        'Crear nuevos usuarios',
        'Asignar roles y permisos',
        'Vincular usuarios a entidades',
        'Activar/desactivar usuarios',
        'Restablecer contraseñas'
    ])

    agregar_subseccion(doc, '3.4 Reportes')
    doc.add_paragraph('Generación de informes:')
    agregar_lista(doc, [
        'Reportes de cumplimiento por entidad',
        'Estadísticas generales del sistema',
        'Exportación de datos'
    ])

    agregar_subseccion(doc, '3.5 Administración Django')
    doc.add_paragraph(
        'Acceso al panel de administración de Django para configuraciones avanzadas '
        'del sistema (uso técnico).'
    )

    # Flujo de trabajo
    agregar_seccion(doc, '4. Flujo de Trabajo Recomendado')
    doc.add_paragraph('Para registrar una nueva entidad en el sistema:')
    agregar_lista(doc, [
        '1. Ingresar al módulo "Entidades"',
        '2. Hacer clic en "Nueva Entidad"',
        '3. Completar datos de la entidad (NIT, razón social, etc.)',
        '4. El sistema creará automáticamente la sede principal',
        '5. Se generará un usuario administrador para la entidad',
        '6. Entregar credenciales al responsable de la entidad'
    ])

    # Módulos próximos
    agregar_modulos_proximos(doc)

    # Soporte
    agregar_seccion(doc, '5. Soporte Técnico')
    doc.add_paragraph('Para soporte técnico o consultas sobre el sistema, contactar al equipo de desarrollo.')

    # Guardar
    doc.save('MANUAL_ADMINISTRADOR.docx')
    print('[OK] Manual del Administrador creado: MANUAL_ADMINISTRADOR.docx')

# ============================================================
# MANUAL CLÍNICA
# ============================================================
def crear_manual_clinica():
    doc = Document()

    # Título
    agregar_titulo_principal(doc, 'MANUAL DE USUARIO\nENTIDAD PRESTADORA')

    doc.add_paragraph()
    p = doc.add_paragraph('Sistema de Habilitación de Servicios de Salud')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph('Resolución 3100 de 2019')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    # Credenciales
    agregar_seccion(doc, '1. Credenciales de Acceso')
    doc.add_paragraph('Utilice las siguientes credenciales para ingresar al sistema:')
    doc.add_paragraph()
    agregar_tabla_credenciales(doc, 'clinica@habilitacion.com', 'Clinica123!', 'Administrador de Entidad')
    doc.add_paragraph()
    doc.add_paragraph('URL de acceso: http://127.0.0.1:8012/')
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('Entidad: ')
    run.bold = True
    p.add_run('CLINICA EJEMPLO SALUD S.A.S.')

    p = doc.add_paragraph()
    run = p.add_run('Código REPS: ')
    run.bold = True
    p.add_run('50001234567')

    # Descripción
    agregar_seccion(doc, '2. Descripción del Sistema')
    doc.add_paragraph(
        'El Sistema de Habilitación de Servicios de Salud le permite gestionar el '
        'proceso de autoevaluación de su entidad según los estándares establecidos '
        'en la Resolución 3100 de 2019 del Ministerio de Salud y Protección Social.'
    )

    doc.add_paragraph(
        'El sistema está organizado por sedes, permitiendo evaluar cada sede de su '
        'entidad de forma independiente, siguiendo la estructura de grupos de estándares:'
    )

    agregar_lista(doc, [
        '11.1 - Estándares aplicables a todos los servicios (OBLIGATORIO)',
        '11.2 - Grupo Consulta Externa',
        '11.3 - Grupo Apoyo Diagnóstico y Complementación Terapéutica',
        '11.4 - Grupo Internación',
        '11.5 - Grupo Quirúrgico',
        '11.6 - Grupo Atención Inmediata'
    ])

    # Funcionalidades
    agregar_seccion(doc, '3. Funcionalidades Disponibles')

    agregar_subseccion(doc, '3.1 Dashboard')
    doc.add_paragraph('Panel principal con estadísticas de su entidad:')
    agregar_lista(doc, [
        'Resumen de cumplimiento por sede',
        'Progreso de evaluación',
        'Criterios pendientes',
        'Indicadores de cumplimiento'
    ])

    agregar_subseccion(doc, '3.2 Evaluación por Sedes')
    doc.add_paragraph('Módulo principal para realizar la autoevaluación:')
    agregar_lista(doc, [
        'Seleccionar la sede a evaluar',
        'Navegar por categorías (11.1, 11.2, etc.)',
        'Navegar por subcategorías (Talento Humano, Infraestructura, etc.)',
        'Evaluar cada criterio: Cumple (C), No Cumple (NC), No Aplica (NA)',
        'Agregar comentarios y justificaciones',
        'Subir documentos de soporte',
        'Ver resumen de evaluación "Evaluar lo que llevo"'
    ])

    agregar_subseccion(doc, '3.3 Configuración de Sede')
    doc.add_paragraph('Habilitar los grupos de estándares que aplican a cada sede:')
    agregar_lista(doc, [
        'El grupo 11.1 es obligatorio y no se puede desactivar',
        'Activar/desactivar grupos 11.2 a 11.6 según servicios ofrecidos',
        'Solo se evaluarán los grupos activos'
    ])

    agregar_subseccion(doc, '3.4 Gestión de Documentos')
    doc.add_paragraph('Para cada criterio puede:')
    agregar_lista(doc, [
        'Subir documentos de soporte (PDF, Word, imágenes)',
        'Ver documentos ya cargados',
        'Eliminar documentos',
        'Los documentos quedan vinculados al criterio específico'
    ])

    agregar_subseccion(doc, '3.5 Información de Entidad')
    doc.add_paragraph('Visualizar y actualizar datos de su entidad:')
    agregar_lista(doc, [
        'Datos generales (razón social, NIT, etc.)',
        'Información de sedes',
        'Representante legal',
        'Datos de contacto'
    ])

    # Flujo de evaluación
    agregar_seccion(doc, '4. Flujo de Evaluación Recomendado')

    agregar_subseccion(doc, 'Paso 1: Configurar Sede')
    agregar_lista(doc, [
        'Ir a Evaluación → Sedes',
        'Seleccionar la sede',
        'Ir a "Configuración"',
        'Activar los grupos de estándares que aplican a su sede'
    ])

    agregar_subseccion(doc, 'Paso 2: Evaluar Criterios')
    agregar_lista(doc, [
        'Seleccionar categoría (ej: 11.1)',
        'Seleccionar subcategoría (ej: Talento Humano)',
        'Para cada criterio, seleccionar estado:',
        '   • C (Cumple): El criterio se cumple completamente',
        '   • NC (No Cumple): El criterio no se cumple',
        '   • NA (No Aplica): El criterio no aplica a su servicio',
        'Agregar comentarios si es necesario',
        'Subir documentos de soporte'
    ])

    agregar_subseccion(doc, 'Paso 3: Verificar Progreso')
    agregar_lista(doc, [
        'Usar botón "Evaluar lo que llevo" para ver resumen',
        'Revisar criterios pendientes',
        'Verificar porcentaje de cumplimiento'
    ])

    # Tipos de criterios
    agregar_seccion(doc, '5. Interpretación de la Pantalla de Evaluación')
    doc.add_paragraph('En la pantalla de criterios verá diferentes tipos de filas:')

    agregar_lista(doc, [
        'Títulos (fondo azul oscuro): Secciones principales, no se evalúan',
        'Subtítulos (fondo gris): Subsecciones, no se evalúan',
        'Criterios (fondo blanco): Son los ítems a evaluar con C/NC/NA'
    ])

    doc.add_paragraph()
    doc.add_paragraph('Los colores de estado son:')
    agregar_lista(doc, [
        'Verde: Cumple (C)',
        'Rojo: No Cumple (NC)',
        'Gris: No Aplica (NA)',
        'Amarillo: Pendiente de evaluar'
    ])

    # Módulos próximos
    agregar_modulos_proximos(doc)

    # Soporte
    agregar_seccion(doc, '6. Soporte')
    doc.add_paragraph(
        'Para dudas sobre el uso del sistema o problemas técnicos, '
        'contacte al administrador del sistema.'
    )

    # Guardar
    doc.save('MANUAL_CLINICA.docx')
    print('[OK] Manual de Clinica creado: MANUAL_CLINICA.docx')

# ============================================================
# MANUAL AUDITOR
# ============================================================
def crear_manual_auditor():
    doc = Document()

    # Titulo
    agregar_titulo_principal(doc, 'MANUAL DE USUARIO\nAUDITOR')

    doc.add_paragraph()
    p = doc.add_paragraph('Sistema de Habilitacion de Servicios de Salud')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph('Resolucion 3100 de 2019')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    # Credenciales
    agregar_seccion(doc, '1. Credenciales de Acceso')
    doc.add_paragraph('Utilice las siguientes credenciales para ingresar al sistema:')
    doc.add_paragraph()
    agregar_tabla_credenciales(doc, 'auditor@habilitacion.com', 'Auditor123!', 'Auditor (Solo lectura)')
    doc.add_paragraph()
    doc.add_paragraph('URL de acceso: http://127.0.0.1:8012/')
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('Entidad asignada: ')
    run.bold = True
    p.add_run('CLINICA EJEMPLO SALUD S.A.S.')

    # Descripcion del rol
    agregar_seccion(doc, '2. Descripcion del Rol de Auditor')
    doc.add_paragraph(
        'Como Auditor del sistema, usted tiene acceso de SOLO LECTURA a toda la '
        'informacion de la entidad asignada. Este rol esta disenado para:'
    )
    agregar_lista(doc, [
        'Auditores externos que necesitan revisar el estado de cumplimiento',
        'Supervisores de secretarias de salud',
        'Consultores de calidad',
        'Personal de entes de control'
    ])

    agregar_nota_importante(doc,
        'El rol de Auditor NO puede modificar ninguna informacion. '
        'Solo puede visualizar evaluaciones, documentos y reportes.'
    )

    # Funcionalidades
    agregar_seccion(doc, '3. Funcionalidades Disponibles')

    agregar_subseccion(doc, '3.1 Dashboard')
    doc.add_paragraph('Panel con estadisticas de la entidad:')
    agregar_lista(doc, [
        'Resumen de cumplimiento por sede',
        'Progreso de evaluacion',
        'Indicadores de cumplimiento'
    ])

    agregar_subseccion(doc, '3.2 Ver Evaluaciones')
    doc.add_paragraph('Puede navegar y ver todas las evaluaciones:')
    agregar_lista(doc, [
        'Ver sedes de la entidad',
        'Ver categorias y subcategorias',
        'Ver estado de cada criterio (Cumple, No Cumple, No Aplica)',
        'Ver comentarios registrados',
        'Ver documentos de soporte cargados'
    ])

    agregar_subseccion(doc, '3.3 Ver Documentos')
    doc.add_paragraph('Acceso a documentos de soporte:')
    agregar_lista(doc, [
        'Visualizar documentos cargados por criterio',
        'Descargar documentos para revision',
        'Ver historial de documentos'
    ])

    agregar_subseccion(doc, '3.4 Ver Reportes')
    doc.add_paragraph('Acceso a reportes de cumplimiento:')
    agregar_lista(doc, [
        'Reporte general de cumplimiento',
        'Estadisticas por estandar',
        'Comparativos entre sedes'
    ])

    # Lo que NO puede hacer
    agregar_seccion(doc, '4. Restricciones del Rol')
    doc.add_paragraph('Como Auditor, usted NO puede:')
    agregar_lista(doc, [
        'Modificar el estado de los criterios',
        'Agregar o eliminar comentarios',
        'Subir o eliminar documentos',
        'Modificar configuraciones de sedes',
        'Editar informacion de la entidad',
        'Crear usuarios'
    ])

    # Modulos proximos
    agregar_modulos_proximos(doc)

    # Soporte
    agregar_seccion(doc, '5. Soporte')
    doc.add_paragraph(
        'Para dudas sobre el uso del sistema, contacte al administrador de la entidad.'
    )

    # Guardar
    doc.save('MANUAL_AUDITOR.docx')
    print('[OK] Manual del Auditor creado: MANUAL_AUDITOR.docx')


# Ejecutar
if __name__ == '__main__':
    print('Creando manuales de usuario...')
    print()
    crear_manual_admin()
    crear_manual_clinica()
    crear_manual_auditor()
    print()
    print('Manuales creados exitosamente!')
